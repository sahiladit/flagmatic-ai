from fastapi import FastAPI, UploadFile, Form, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
import google.generativeai as genai
import os
from dotenv import load_dotenv
import pdfplumber
import tempfile
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement



# Load API key from .env
load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

app = FastAPI()

from fastapi.staticfiles import StaticFiles
app.mount("/", StaticFiles(directory="static", html=True), name="static")

from fastapi.responses import FileResponse
from pathlib import Path

@app.get("/")
async def root():
    return FileResponse(Path("static/index.html"))


# Enable CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class ResumeJobRequest(BaseModel):
    resume: str
    job: str

class TailorAIAgent:
    def __init__(self, model):
        self.model = model

    def generate_tailored_resume(self, resume, job):
        prompt = f"""
        Rewrite the following resume to better fit the job description. Keep it professional, ATS-friendly and impactful.
        Resume:
        {resume}

        Job Description:
        {job}
        """
        response = self.model.generate_content(prompt)
        return response.text.strip()

def save_resume_to_docx(text):
    doc = Document()
    doc.add_heading("Tailored Resume", level=1)

    paragraphs = text.split("\n")
    for para in paragraphs:
        if "**" in para:
            run_parts = para.split("**")
            p = doc.add_paragraph()
            for i, part in enumerate(run_parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
        else:
            doc.add_paragraph(para)

    temp_doc_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(temp_doc_path)
    return temp_doc_path

@app.post("/agent-generate")
async def agent_generate(req: ResumeJobRequest):
    model = genai.GenerativeModel("gemini-2.0-flash")
    agent = TailorAIAgent(model)
    tailored_resume = agent.generate_tailored_resume(req.resume, req.job)
    doc_path = save_resume_to_docx(tailored_resume)
    return FileResponse(path=doc_path, filename="Tailored_Resume.docx", media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.post("/agent-upload")
async def agent_upload(file: UploadFile = File(...), job: str = Form(...)):
    with tempfile.NamedTemporaryFile(delete=False) as tmp:
        contents = await file.read()
        tmp.write(contents)
        tmp_path = tmp.name

    resume_text = ""
    try:
        with pdfplumber.open(tmp_path) as pdf:
            for page in pdf.pages:
                resume_text += page.extract_text() + "\n"
    except Exception:
        return {"error": "Failed to parse PDF resume."}

    model = genai.GenerativeModel("gemini-2.0-flash")
    agent = TailorAIAgent(model)
    tailored_resume = agent.generate_tailored_resume(resume_text, job)
    doc_path = save_resume_to_docx(tailored_resume)
    return FileResponse(path=doc_path, filename="Tailored_Resume.docx", media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')